from typing import List, Sequence, Mapping, Any
from jinja2 import Template
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from io import BytesIO
from .diagram_analysis import Features
from .process_breakdown import ProcessStep
from .company_matching import Match

# optional Word(.docx) support
try:
  from docx import Document  # type: ignore
except Exception:  # pragma: no cover
  Document = None  # type: ignore

_TEMPLATE = Template(
    """
    <div class="report">
      <h2>CMA マッチングレポート</h2>
      <section>
        <h3>1. 図面解析結果</h3>
        <ul>
          <li>ファイル: {{ f.filename }} ({{ f.ext }})</li>
          <li>材質候補: {{ f.material or '不明' }}</li>
          <li>部品種別候補: {{ f.part_type or '不明' }}</li>
          {% if f.title %}<li>タイトル: {{ f.title }}</li>{% endif %}
          {% if f.drawing_no %}<li>図番: {{ f.drawing_no }}</li>{% endif %}
          {% if f.surface_finish %}<li>表面粗さ: {{ f.surface_finish }}</li>{% endif %}
          {% if f.tolerances %}<li>公差: {{ ', '.join(f.tolerances) }}</li>{% endif %}
          {% if f.recommended_process or f.recommended_machine %}
            <li>推奨加工/装置: {{ f.recommended_process or '-' }} / {{ f.recommended_machine or '-' }}</li>
          {% endif %}
        </ul>
        {% if f.notes %}
        <details><summary>抽出テキスト</summary><pre>{{ f.notes }}</pre></details>
        {% endif %}
      </section>

      <section>
        <h3>2. 加工工程案</h3>
        <table border="1" cellspacing="0" cellpadding="6">
          <tr><th>工程</th><th>装置</th><th>目安時間(min)</th><th>公差</th><th>精度</th></tr>
          {% for s in steps %}
          <tr>
            <td>{{ s.name }}</td>
            <td>{{ s.machine }}</td>
            <td>{{ s.minutes }}</td>
            <td>{{ s.tolerance or '-' }}</td>
            <td>{{ s.precision or '-' }}</td>
          </tr>
          {% endfor %}
        </table>
      </section>

      <section>
        <h3>3. 企業マッチング結果</h3>
        <table border="1" cellspacing="0" cellpadding="6">
          <tr><th>企業</th><th>スコア</th><th>対応工程</th></tr>
          {% for m in matches %}
          <tr>
            <td>{{ m.company.name }}</td>
            <td>{{ '%.2f' % m.score }}</td>
            <td>{{ ', '.join(m.steps) if m.steps else '-' }}</td>
          </tr>
          {% endfor %}
        </table>
        {% if matches and matches[0].alliance %}
        <p><strong>アライアンス提案:</strong>
          {% for c in matches[0].alliance %}
            {{ c.name }}{% if not loop.last %}, {% endif %}
          {% endfor %}
        </p>
        {% endif %}
      </section>
    </div>
    """
)


def render_report_html(f: Features, steps: List[ProcessStep], matches: List[Match]) -> str:
  return _TEMPLATE.render(f=f, steps=steps, matches=matches)


def render_report_pdf(f: Features, steps: List[ProcessStep], matches: List[Match]) -> bytes:
  buf = BytesIO()
  c = canvas.Canvas(buf, pagesize=A4)
  width, height = A4
  y = height - 20 * mm
  c.setFont("Helvetica-Bold", 16)
  c.drawString(20 * mm, y, "CMA マッチングレポート")
  y -= 12 * mm
  c.setFont("Helvetica", 10)
  # 1. Features
  c.drawString(20 * mm, y, f"ファイル: {f.filename} ({f.ext})  材質: {f.material or '不明'}  種別: {f.part_type or '不明'}")
  y -= 8 * mm
  # 2. Steps
  c.setFont("Helvetica-Bold", 12)
  c.drawString(20 * mm, y, "加工工程案")
  y -= 7 * mm
  c.setFont("Helvetica", 10)
  for s in steps:
    line = f"- {s.name} / {s.machine} / {s.minutes}min / tol: {s.tolerance or '-'}"
    c.drawString(22 * mm, y, line)
    y -= 6 * mm
    if y < 20 * mm:
      c.showPage(); y = height - 20 * mm
  # 3. Matches
  c.setFont("Helvetica-Bold", 12)
  c.drawString(20 * mm, y, "企業マッチング結果")
  y -= 7 * mm
  c.setFont("Helvetica", 10)
  for m in matches:
    line = f"- {m.company.name}  score={m.score}  steps: {', '.join(m.steps) if m.steps else '-'}"
    c.drawString(22 * mm, y, line)
    y -= 6 * mm
    if y < 20 * mm:
      c.showPage(); y = height - 20 * mm
  c.showPage()
  c.save()
  return buf.getvalue()


def render_report_docx(f: Features, steps: List[ProcessStep], matches: List[Match]) -> bytes:
  """マッチングレポートを Word(.docx) として生成する"""
  if Document is None:
    raise RuntimeError("python-docx がインストールされていません。requirements.txt を更新してインストールしてください。")

  doc = Document()
  doc.add_heading('CMA マッチングレポート', level=1)

  # 1. 図面解析結果
  doc.add_heading('1. 図面解析結果', level=2)
  doc.add_paragraph(f"ファイル: {f.filename} ({f.ext})")
  doc.add_paragraph(f"材質候補: {f.material or '不明'}")
  doc.add_paragraph(f"部品種別候補: {f.part_type or '不明'}")
  if getattr(f, 'title', None):
    doc.add_paragraph(f"タイトル: {getattr(f, 'title')}")
  if getattr(f, 'drawing_no', None):
    doc.add_paragraph(f"図番: {getattr(f, 'drawing_no')}")
  if getattr(f, 'surface_finish', None):
    doc.add_paragraph(f"表面粗さ: {getattr(f, 'surface_finish')}")
  tol = getattr(f, 'tolerances', None)
  if tol:
    doc.add_paragraph("公差: " + ", ".join(tol))
  rp = getattr(f, 'recommended_process', None)
  rm = getattr(f, 'recommended_machine', None)
  if rp or rm:
    doc.add_paragraph(f"推奨加工/装置: {rp or '-'} / {rm or '-'}")
  if getattr(f, 'notes', None):
    doc.add_paragraph('抽出テキスト:')
    doc.add_paragraph(getattr(f, 'notes'))

  # 2. 加工工程案
  doc.add_heading('2. 加工工程案', level=2)
  if steps:
    tbl = doc.add_table(rows=1, cols=5)
    hdr = tbl.rows[0].cells
    hdr[0].text = '工程'
    hdr[1].text = '装置'
    hdr[2].text = '目安時間(min)'
    hdr[3].text = '公差'
    hdr[4].text = '精度'
    for s in steps:
      row = tbl.add_row().cells
      row[0].text = s.name or ''
      row[1].text = s.machine or ''
      row[2].text = str(s.minutes)
      row[3].text = s.tolerance or '-'
      row[4].text = getattr(s, 'precision', None) or '-'
  else:
    doc.add_paragraph('工程情報なし')

  # 3. 企業マッチング結果
  doc.add_heading('3. 企業マッチング結果', level=2)
  if matches:
    tbl = doc.add_table(rows=1, cols=4)
    hdr = tbl.rows[0].cells
    hdr[0].text = '企業'
    hdr[1].text = 'スコア'
    hdr[2].text = '対応工程'
    hdr[3].text = 'アライアンス提案'
    for m in matches:
      row = tbl.add_row().cells
      row[0].text = m.company.name
      row[1].text = f"{m.score:.2f}"
      row[2].text = ', '.join(m.steps) if m.steps else '-'
    # アライアンス提案
    if getattr(matches[0], 'alliance', None):
      doc.add_paragraph('アライアンス提案: ' + ', '.join([c.name for c in matches[0].alliance]))
  else:
    doc.add_paragraph('候補なし')

  out = BytesIO()
  doc.save(out)
  return out.getvalue()


def render_assignments_docx(drawing_file: str, items: Sequence[Mapping[str, Any]]) -> bytes:
  """選択された図面に対する割当一覧のみを Word(.docx) で出力する簡易レポート"""
  if Document is None:
    raise RuntimeError("python-docx がインストールされていません。requirements.txt を更新してインストールしてください。")

  doc = Document()
  doc.add_heading('CMA 割当レポート', level=1)
  doc.add_paragraph(f"図面: {drawing_file}")
  doc.add_paragraph("この文書は選択された図面に対する企業へのタスク割当のみを含みます。")

  tbl = doc.add_table(rows=1, cols=4)
  hdr = tbl.rows[0].cells
  hdr[0].text = 'ID'
  hdr[1].text = 'Task'
  hdr[2].text = 'Company'
  hdr[3].text = 'Created'
  for it in items:
    row = tbl.add_row().cells
    row[0].text = str(it.get('id', ''))
    row[1].text = str(it.get('task_name', ''))
    row[2].text = str(it.get('company_name', ''))
    row[3].text = str(it.get('created_at', ''))

  out = BytesIO()
  doc.save(out)
  return out.getvalue()
